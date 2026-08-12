import io
import math
import logging
from pathlib import Path
from typing import Optional, List, Tuple
from PIL import Image

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    Document = None


def generate_student_grid_report(
    student_photos: List[Tuple[str, Path]],
    output_docx_path: Path | str,
    logger: Optional[logging.Logger] = None,
) -> bool:
    """Gather individual student images and generate a single-page Word report in grid format."""
    if Document is None:
        if logger:
            logger.warning("python-docx is not installed. Skipping Word grid report generation.")
        return False

    if not student_photos:
        if logger:
            logger.warning("No student photos available to generate Word report.")
        return False

    doc = Document()

    # Narrow margins (0.4 in) to fit all elements onto a single page
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

    # Header title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(8)
    run = title_p.add_run("Student Reference Roster Report")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    n_students = len(student_photos)

    # Adjust grid layout & image sizes dynamically for single-page layout
    if n_students <= 6:
        cols = 3
        img_size = Inches(1.8)
    elif n_students <= 12:
        cols = 4
        img_size = Inches(1.3)
    elif n_students <= 20:
        cols = 5
        img_size = Inches(1.0)
    else:
        cols = 6
        img_size = Inches(0.8)

    rows = math.ceil(n_students / cols)
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, (student_name, img_path) in enumerate(student_photos):
        r = idx // cols
        c = idx % cols
        cell = table.cell(r, c)

        cell_p = cell.paragraphs[0]
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_p.paragraph_format.space_before = Pt(2)
        cell_p.paragraph_format.space_after = Pt(2)

        # Convert image (including .webp) to PNG stream in-memory
        try:
            with Image.open(img_path) as img:
                img_stream = io.BytesIO()
                img.save(img_stream, format="PNG")
                img_stream.seek(0)

                run_img = cell_p.add_run()
                run_img.add_picture(img_stream, width=img_size)
        except Exception as e:
            if logger:
                logger.error(f"Error adding image {img_path.name} to Word report: {e}")

        # Add student caption label
        p_name = cell.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_before = Pt(1)
        p_name.paragraph_format.space_after = Pt(4)
        run_name = p_name.add_run(student_name)
        run_name.font.name = "Calibri"
        run_name.font.size = Pt(9)
        run_name.font.bold = True

    output_path = Path(output_docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    if logger:
        logger.info(f"Generated single-page student roster report: {output_path}")
    return True
